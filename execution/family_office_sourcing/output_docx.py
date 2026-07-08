"""Renders the prospect list as a house-style Word document. Reuses dd_agent's docx
styling helpers directly (same navy/gold/cream/Cambria palette, same native-table
approach so it pastes cleanly into other firm documents) rather than re-implementing them —
both modules live on the same execution/ sys.path, and duplicating ~100 lines of Word OOXML
plumbing for an identical house style would be the kind of premature-abstraction-avoidance
that actually just risks the two report styles drifting apart.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches

from dd_agent.config import ReportConfig
from dd_agent.reporting.docx_report import (
    _add_bullets,
    _add_heading,
    _add_paragraph,
    _add_table,
    _hex_to_rgb,
    _set_title_rule_color,
    _style_base_fonts,
)
from family_office_sourcing.schemas import CandidateRecord, Confidence, PriorityTier, SourcingRun

_TIER_LABELS = {
    PriorityTier.A: "Tier A — Priority Outreach",
    PriorityTier.B: "Tier B — Qualified, Lower Priority",
    PriorityTier.C: "Tier C — Unconfirmed / Low Fit",
}


_ROW_HEADERS = ["Name", "City", "State", "Principal", "Sector Signals", "Website", "Confidence", "Mode", "Source(s)", "Note"]
_ROW_COLUMN_WIDTHS = [1.3, 0.6, 0.6, 1.0, 1.0, 1.2, 0.6, 0.9, 1.3, 1.5]


def _row(record: CandidateRecord) -> list[str]:
    sources = "; ".join(sorted({s.source_name for s in record.sources}))
    # An exclusion (non-family-office pattern match) is a more important thing to surface
    # in this column than an enrichment rationale — it explains why the record is pinned to
    # Tier C in a way "plausible" alone wouldn't.
    note = record.exclusion_reason or (record.enrichment.rationale if record.enrichment else "—")
    return [
        record.display_name,
        record.city or "—",
        record.state or "—",
        record.principal or "—",
        ", ".join(record.sector_signals) or "—",
        record.website or "—",
        record.confidence.value,
        record.investment_mode.value,
        sources,
        note,
    ]


def _render_records_table(document: Document, records: list[CandidateRecord], report_config: ReportConfig) -> None:
    if not records:
        _add_paragraph(document, "No candidates in this group for this run.", report_config)
        return
    rows = [_row(r) for r in sorted(records, key=lambda r: -r.priority_score)]
    _add_table(document, _ROW_HEADERS, rows, report_config, column_widths_inches=_ROW_COLUMN_WIDTHS)


def _render_tier(document: Document, tier: PriorityTier, records: list[CandidateRecord], report_config: ReportConfig) -> None:
    _add_heading(document, _TIER_LABELS[tier], 1, report_config)
    _render_records_table(document, records, report_config)


def _set_landscape(document: Document) -> None:
    """8-column prospect table doesn't fit portrait's ~6.5in usable width — widen the page
    instead of cramming columns, consistent with dd_agent's "no cramped/wrapped columns"
    fix in its own docx renderer."""
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = section.right_margin = Inches(0.5)


def render_docx_report(run: SourcingRun, report_config: ReportConfig, output_path: Path) -> Path:
    document = Document()
    _set_landscape(document)
    _style_base_fonts(document, report_config)

    title = _add_heading(document, "Indian Family Office Prospects", 0, report_config)
    _set_title_rule_color(title, report_config.palette.gold)
    _add_paragraph(document, f"Run: {run.run_id}", report_config, italic=True)
    _add_paragraph(document, f"Generated: {run.generated_at.isoformat()}", report_config, italic=True)

    by_tier: dict[PriorityTier, list[CandidateRecord]] = {t: [] for t in (PriorityTier.A, PriorityTier.B, PriorityTier.C)}
    for record in run.records:
        by_tier.setdefault(record.priority_tier, []).append(record)

    confirmed = sum(1 for r in run.records if r.confidence == Confidence.CONFIRMED)
    unconfirmed = len(run.records) - confirmed
    _add_bullets(
        document,
        [
            f"{len(run.records)} candidates this run — {confirmed} confirmed (structured database), {unconfirmed} unconfirmed (signal-derived, needs verification).",
            f"Tier A: {len(by_tier.get(PriorityTier.A, []))} | Tier B: {len(by_tier.get(PriorityTier.B, []))} | Tier C: {len(by_tier.get(PriorityTier.C, []))}",
            "Unconfirmed candidates and any row missing a Website/Source should be independently verified before outreach — see the directive's non-goals on contact data.",
        ],
        report_config,
    )

    for tier in (PriorityTier.A, PriorityTier.B, PriorityTier.C):
        _render_tier(document, tier, by_tier.get(tier, []), report_config)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path


_NO_STATE_LABEL = "No State on File"


def render_docx_report_by_state(run: SourcingRun, report_config: ReportConfig, output_path: Path) -> Path:
    """Same data and row shape as render_docx_report, grouped by state instead of tier —
    one all-India section per state instead of dozens of near-empty per-state files (most
    states have zero or one candidate in any given run; grouping keeps this to one file)."""
    document = Document()
    _set_landscape(document)
    _style_base_fonts(document, report_config)

    title = _add_heading(document, "Indian Family Office Prospects — By State", 0, report_config)
    _set_title_rule_color(title, report_config.palette.gold)
    _add_paragraph(document, f"Run: {run.run_id}", report_config, italic=True)
    _add_paragraph(document, f"Generated: {run.generated_at.isoformat()}", report_config, italic=True)

    by_state: dict[str, list[CandidateRecord]] = {}
    for record in run.records:
        by_state.setdefault(record.state or _NO_STATE_LABEL, []).append(record)

    ordered_states = sorted((s for s in by_state if s != _NO_STATE_LABEL), key=lambda s: -len(by_state[s]))
    if _NO_STATE_LABEL in by_state:
        ordered_states.append(_NO_STATE_LABEL)

    confirmed = sum(1 for r in run.records if r.confidence == Confidence.CONFIRMED)
    unconfirmed = len(run.records) - confirmed
    state_breakdown = ", ".join(f"{s}: {len(by_state[s])}" for s in ordered_states if s != _NO_STATE_LABEL)
    _add_bullets(
        document,
        [
            f"{len(run.records)} candidates this run across {len(ordered_states) - (_NO_STATE_LABEL in by_state)} states — "
            f"{confirmed} confirmed (structured database), {unconfirmed} unconfirmed (signal-derived, needs verification).",
            f"By state: {state_breakdown}." if state_breakdown else "No state data available in this run.",
            "Unconfirmed candidates and any row missing a Website/Source should be independently verified before outreach — see the directive's non-goals on contact data.",
        ],
        report_config,
    )

    for state in ordered_states:
        _add_heading(document, f"{state} ({len(by_state[state])})", 1, report_config)
        _render_records_table(document, by_state[state], report_config)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path
