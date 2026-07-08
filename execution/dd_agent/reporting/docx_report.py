"""Renders the DD summary report as a formatted Word document, per house style.

Same data and section structure as markdown_report.py (both render
dd_agent.reporting.data.ReportContext) — this module only adds styling: Cambria
headings/body, restrained navy/gold/cream palette (config.yaml `report`), native Word
Heading styles and native table objects (not images/text blocks) so the financial summary
and risk register can be selected and pasted directly into the firm's existing DD report
template.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from dd_agent.config import ReportConfig
from dd_agent.reporting.data import ReportContext, build_report_context
from dd_agent.schemas.contracts import ContractsRegister
from dd_agent.schemas.document_index import DocumentIndex
from dd_agent.schemas.financial import FinancialSummary
from dd_agent.schemas.risk_register import RiskRegister


def _resolve_color(value: str, report_config: ReportConfig) -> str:
    """`value` may be a palette name (navy/gold/cream, per config.yaml `report.table`) or a
    literal hex code — resolve the former against report_config.palette, pass through the
    latter unchanged."""
    named = {
        "navy": report_config.palette.navy,
        "gold": report_config.palette.gold,
        "cream": report_config.palette.cream,
    }
    return named.get(value.lower(), value)


def _hex_to_rgb(hex_str: str) -> RGBColor:
    return RGBColor.from_string(hex_str.lstrip("#").upper())


def _set_run_font(run, font_name: str, size: float | None = None, color: RGBColor | None = None, bold: bool = False) -> None:
    run.font.name = font_name
    # Word falls back to a default east-asian font for some runs unless this is set too,
    # which can silently defeat a heading-font choice like Cambria in some Word versions.
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = r_pr.makeelement(qn("w:rFonts"), {})
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    run.font.bold = bold


def _set_title_rule_color(paragraph, hex_color: str) -> None:
    """Word's built-in "Title" paragraph style ships with a bottom border in its own
    default blue, unrelated to our palette — recolor it instead of leaving stock Word blue
    on the one document title. This is also the only place gold (the house-style accent)
    actually gets used, per "gold is a sparing accent — rules, small header touches"."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = p_pr.makeelement(qn("w:pBdr"), {})
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = p_bdr.makeelement(qn("w:bottom"), {})
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), hex_color.lstrip("#"))


def _set_column_widths(table, widths_inches: list[float]) -> None:
    """python-docx tables default to splitting width evenly across columns regardless of
    content, which produces cramped/wrapped short columns next to starved long ones (seen
    directly when the rendered docx was exported to PDF and inspected). Word requires the
    width set on both the table grid and every individual cell for fixed-layout tables to
    render reliably."""
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.makeelement(qn("w:tblLayout"), {qn("w:type"): "fixed"})
    tbl_pr.append(layout)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_inches[idx])
    for idx, column in enumerate(table.columns):
        column.width = Inches(widths_inches[idx])


def _shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.makeelement(
        qn("w:shd"), {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color.lstrip("#")}
    )
    tc_pr.append(shd)


def _style_base_fonts(document: Document, report_config: ReportConfig) -> None:
    normal = document.styles["Normal"]
    normal.font.name = report_config.body_font
    normal.font.size = Pt(10.5)
    for level in range(1, 4):
        style = document.styles[f"Heading {level}"]
        style.font.name = report_config.heading_font
        style.font.color.rgb = _hex_to_rgb(report_config.palette.navy)
        style.font.bold = True


def _add_heading(document: Document, text: str, level: int, report_config: ReportConfig):
    heading = document.add_heading(text, level=level)
    for run in heading.runs:
        _set_run_font(run, report_config.heading_font, color=_hex_to_rgb(report_config.palette.navy), bold=True)
    return heading


def _add_paragraph(document: Document, text: str, report_config: ReportConfig, italic: bool = False, bold_prefix: str | None = None):
    paragraph = document.add_paragraph()
    if bold_prefix:
        prefix_run = paragraph.add_run(bold_prefix)
        _set_run_font(prefix_run, report_config.body_font, size=10.5, color=_hex_to_rgb(report_config.palette.navy), bold=True)
    run = paragraph.add_run(text)
    run.italic = italic
    _set_run_font(run, report_config.body_font, size=10.5, color=_hex_to_rgb(report_config.palette.navy))
    return paragraph


def _add_bullets(document: Document, items: list[str], report_config: ReportConfig) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        _set_run_font(run, report_config.body_font, size=10.5, color=_hex_to_rgb(report_config.palette.navy))


def _add_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    report_config: ReportConfig,
    column_widths_inches: list[float] | None = None,
):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    header_text_color = _hex_to_rgb(_resolve_color(report_config.table.header_text, report_config))
    header_fill_color = _resolve_color(report_config.table.header_fill, report_config)
    body_text_color = _hex_to_rgb(_resolve_color(report_config.table.body_text, report_config))

    header_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        header_cells[i].text = ""
        run = header_cells[i].paragraphs[0].add_run(header_text)
        _set_run_font(run, report_config.body_font, size=10, color=header_text_color, bold=True)
        _shade_cell(header_cells[i], header_fill_color)

    for row_values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(value))
            _set_run_font(run, report_config.body_font, size=10, color=body_text_color)

    if column_widths_inches:
        _set_column_widths(table, column_widths_inches)

    return table


def _render_executive_summary(document: Document, ctx: ReportContext, report_config: ReportConfig) -> None:
    _add_heading(document, "Executive Summary", 1, report_config)
    bullets = [f"{ctx.exec_summary.total_documents} documents reviewed from the data room."]
    if ctx.exec_summary.documents_by_category:
        breakdown = ", ".join(
            f"{v} {k}" for k, v in sorted(ctx.exec_summary.documents_by_category.items(), key=lambda kv: -kv[1])
        )
        bullets.append(f"Document mix: {breakdown}.")
    rc = ctx.exec_summary.risk_counts_by_severity
    bullets.append(
        f"{sum(rc.values())} risk register items identified: "
        f"{rc.get('High', 0)} High, {rc.get('Medium', 0)} Medium, {rc.get('Low', 0)} Low."
    )
    bullets.append(
        f"Financial data covers: {', '.join(ctx.exec_summary.periods_covered)}."
        if ctx.exec_summary.periods_covered
        else "No financial statement data was extracted from the data room."
    )
    bullets.append(f"{ctx.exec_summary.contracts_reviewed} contract(s) reviewed and entered into the contracts register.")
    if ctx.exec_summary.needs_attention_count:
        bullets.append(
            f"{ctx.exec_summary.needs_attention_count} document(s) could not be read "
            f"(locked or unreadable) — see Open Questions for Management."
        )
    if ctx.category_gaps:
        bullets.append(f"No documentation was provided for: {', '.join(ctx.category_gaps)}.")
    _add_bullets(document, bullets, report_config)


def _render_corporate_structure(document: Document, ctx: ReportContext, report_config: ReportConfig) -> None:
    _add_heading(document, "Corporate Structure", 1, report_config)
    if not ctx.corporate_docs:
        _add_paragraph(document, "No Corporate/Legal documents were provided in the data room.", report_config)
        return
    _add_paragraph(
        document,
        "v1 does not perform dedicated corporate structure extraction; the Corporate/Legal "
        "documents identified are listed below for analyst review.",
        report_config,
        italic=True,
    )
    _add_bullets(
        document,
        [f"{doc.file_path} — {doc.summary or '(no summary available)'}" for doc in ctx.corporate_docs],
        report_config,
    )


def _render_financial_overview(document: Document, ctx: ReportContext, report_config: ReportConfig) -> None:
    _add_heading(document, "Financial Overview", 1, report_config)
    if not ctx.financial_periods:
        _add_paragraph(document, "No financial statement data was extracted from the data room.", report_config)
        return

    headers = ["Line Item"] + [p.period_label for p in ctx.financial_periods]
    rows_spec = {
        "Revenue": [p.revenue for p in ctx.financial_periods],
        "EBITDA (Reported)": [p.ebitda_reported for p in ctx.financial_periods],
        "EBITDA (Adjusted)": [p.ebitda_adjusted for p in ctx.financial_periods],
        "Net Working Capital": [p.net_working_capital for p in ctx.financial_periods],
        "Total Debt": [p.total_debt for p in ctx.financial_periods],
        "Related-Party Revenue": [p.related_party_revenue for p in ctx.financial_periods],
    }
    rows = [
        [label] + [f"{v:,.0f}" if v is not None else "—" for v in values]
        for label, values in rows_spec.items()
    ]
    period_col_width = (6.5 - 2.0) / max(len(ctx.financial_periods), 1)
    _add_table(document, headers, rows, report_config, column_widths_inches=[2.0] + [period_col_width] * len(ctx.financial_periods))

    adjustments = sorted({item for p in ctx.financial_periods for item in p.ebitda_adjustments})
    if adjustments:
        _add_paragraph(document, "; ".join(adjustments), report_config, bold_prefix="EBITDA normalization adjustments noted: ")

    sources = sorted({p.source_file for p in ctx.financial_periods if p.source_file})
    if sources:
        _add_paragraph(document, ", ".join(sources), report_config, italic=True, bold_prefix="Source: ")


def _render_key_risks(document: Document, ctx: ReportContext, report_config: ReportConfig) -> None:
    _add_heading(document, "Key Risks & Red Flags", 1, report_config)
    if not ctx.risk_items:
        _add_paragraph(document, "No risk register items were identified.", report_config)
        return
    for item in ctx.risk_items:
        _add_heading(document, f"[{item.id}] {item.title} — {item.severity.value}", 2, report_config)
        _add_paragraph(
            document, f"Category: {item.category.value} | Confidence: {item.confidence.value}", report_config, italic=True
        )
        _add_paragraph(document, item.description, report_config)
        sources = "; ".join(s.file + (f", p.{s.page}" if s.page else "") for s in item.source_documents)
        _add_paragraph(document, sources, report_config, bold_prefix="Source: ")
        _add_paragraph(document, item.recommended_question, report_config, bold_prefix="Recommended question for management: ")


def _render_contracts_summary(document: Document, ctx: ReportContext, report_config: ReportConfig) -> None:
    _add_heading(document, "Contracts Summary", 1, report_config)
    if not ctx.contracts:
        _add_paragraph(document, "No contracts were identified in the Corporate/Legal documents provided.", report_config)
        return
    headers = ["Counterparty", "Type", "Term End", "Value", "Flags", "Source"]
    rows = []
    for c in ctx.contracts:
        flags = "; ".join(f.description for f in c.flags) if c.flags else "—"
        term_end = c.term_end_date.isoformat() if c.term_end_date else "—"
        rows.append([c.counterparty or "—", c.contract_type or "—", term_end, c.contract_value or "—", flags, c.file_path])
    _add_table(document, headers, rows, report_config, column_widths_inches=[1.4, 0.6, 0.75, 0.85, 1.7, 1.2])


def _render_open_questions(document: Document, ctx: ReportContext, report_config: ReportConfig) -> None:
    _add_heading(document, "Open Questions for Management", 1, report_config)
    if not ctx.open_questions:
        _add_paragraph(document, "No open questions identified.", report_config)
        return
    _add_bullets(
        document,
        [f"[{q.category}] {q.question} (source: {q.source_file})" for q in ctx.open_questions],
        report_config,
    )


def _render_document_index(document: Document, ctx: ReportContext, report_config: ReportConfig) -> None:
    _add_heading(document, "Document Index", 1, report_config)
    headers = ["File", "Category", "Status", "Pages"]
    rows = [
        [e.file_path, e.category.value, e.status.value, str(e.page_count) if e.page_count else "—"]
        for e in ctx.document_index_entries
    ]
    _add_table(document, headers, rows, report_config, column_widths_inches=[3.2, 1.7, 0.8, 0.6])


def render_docx_report(
    index: DocumentIndex,
    risk_register: RiskRegister,
    contracts_register: ContractsRegister,
    financial_summary: FinancialSummary,
    report_config: ReportConfig,
    output_path: Path,
) -> Path:
    ctx = build_report_context(index, risk_register, contracts_register, financial_summary)

    document = Document()
    _style_base_fonts(document, report_config)

    title = _add_heading(document, "Due Diligence Summary", 0, report_config)
    _set_title_rule_color(title, report_config.palette.gold)
    _add_paragraph(document, f"Data room: {ctx.dataroom_path}", report_config, italic=True)
    _add_paragraph(document, f"Run: {ctx.run_id}", report_config, italic=True)

    _render_executive_summary(document, ctx, report_config)
    _render_corporate_structure(document, ctx, report_config)
    _render_financial_overview(document, ctx, report_config)
    _render_key_risks(document, ctx, report_config)
    _render_contracts_summary(document, ctx, report_config)
    _render_open_questions(document, ctx, report_config)
    _render_document_index(document, ctx, report_config)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path
