"""Text extraction per file type, with OCR fallback for scanned/image-only PDFs.

Every extractor returns an `ExtractionResult` and never raises for expected failure modes
(locked/corrupt files) — those are captured as a status + error message so one bad file
never aborts the ingestion run (see directives/due_diligence_agent.md edge cases).
"""
from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from dd_agent.schemas.common import DocumentStatus

logger = logging.getLogger("dd_agent.ingestion.extractors")

# Below this many non-whitespace characters per page, treat a PDF page as "scanned"
# (no embedded text layer) and fall back to OCR.
_MIN_CHARS_PER_PAGE_BEFORE_OCR = 20


@dataclass
class PageText:
    page_number: int
    text: str
    ocr_used: bool = False


@dataclass
class ExtractionResult:
    status: DocumentStatus = DocumentStatus.OK
    text: str = ""
    page_count: Optional[int] = None
    ocr_used: bool = False
    ocr_confidence: Optional[float] = None
    error: Optional[str] = None
    pages: list[PageText] = field(default_factory=list)


def _ocr_image(image) -> tuple[str, Optional[float]]:
    """Run pytesseract on a PIL image; returns (text, mean_confidence)."""
    import pytesseract

    text = pytesseract.image_to_string(image)
    data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
    confidences = [int(c) for c in data.get("conf", []) if str(c).lstrip("-").isdigit() and int(c) >= 0]
    mean_confidence = sum(confidences) / len(confidences) if confidences else None
    return text, mean_confidence


def extract_pdf(path: Path, ocr_enabled: bool, ocr_min_confidence: int) -> ExtractionResult:
    try:
        import fitz  # PyMuPDF
    except ImportError as exc:
        return ExtractionResult(status=DocumentStatus.UNREADABLE, error=f"PyMuPDF not installed: {exc}")

    try:
        doc = fitz.open(str(path))
    except Exception as exc:  # corrupt / unreadable file
        logger.warning("Failed to open PDF %s: %s", path, exc)
        return ExtractionResult(status=DocumentStatus.UNREADABLE, error=str(exc))

    if doc.needs_pass:
        doc.close()
        return ExtractionResult(status=DocumentStatus.LOCKED, error="Password-protected PDF")

    pages: list[PageText] = []
    ocr_used_any = False
    ocr_confidences: list[float] = []

    try:
        for page_index in range(doc.page_count):
            page = doc.load_page(page_index)
            text = page.get_text().strip()

            if len(text) < _MIN_CHARS_PER_PAGE_BEFORE_OCR and ocr_enabled:
                try:
                    import pytesseract  # noqa: F401  (import check before rendering)
                    from PIL import Image

                    pix = page.get_pixmap(dpi=300)
                    image = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
                    ocr_text, confidence = _ocr_image(image)
                    if ocr_text.strip():
                        text = ocr_text.strip()
                        ocr_used_any = True
                        if confidence is not None:
                            ocr_confidences.append(confidence)
                except ImportError:
                    logger.warning("pytesseract/Pillow not installed; skipping OCR fallback for %s", path)
                except Exception as exc:
                    # e.g. pytesseract.TesseractNotFoundError when the Tesseract binary
                    # itself isn't installed. Never let one page's OCR failure crash the
                    # whole ingestion run — fall back to whatever embedded text there was.
                    logger.warning("OCR fallback failed for %s page %d: %s", path, page_index + 1, exc)

            pages.append(PageText(page_number=page_index + 1, text=text, ocr_used=ocr_used_any))
    finally:
        doc.close()

    full_text = "\n\n".join(p.text for p in pages)
    mean_ocr_confidence = (
        sum(ocr_confidences) / len(ocr_confidences) if ocr_confidences else None
    )
    status = DocumentStatus.OK
    if ocr_used_any and mean_ocr_confidence is not None and mean_ocr_confidence < ocr_min_confidence:
        status = DocumentStatus.OCR_LOW_CONFIDENCE

    return ExtractionResult(
        status=status,
        text=full_text,
        page_count=len(pages),
        ocr_used=ocr_used_any,
        ocr_confidence=mean_ocr_confidence,
        pages=pages,
    )


def extract_docx(path: Path) -> ExtractionResult:
    try:
        import docx
    except ImportError as exc:
        return ExtractionResult(status=DocumentStatus.UNREADABLE, error=f"python-docx not installed: {exc}")

    try:
        document = docx.Document(str(path))
    except Exception as exc:
        logger.warning("Failed to open DOCX %s: %s", path, exc)
        return ExtractionResult(status=DocumentStatus.UNREADABLE, error=str(exc))

    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    table_text = []
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                table_text.append(row_text)

    full_text = "\n".join(paragraphs + table_text)
    return ExtractionResult(status=DocumentStatus.OK, text=full_text, page_count=None)


def extract_xlsx(path: Path) -> ExtractionResult:
    try:
        import openpyxl
    except ImportError as exc:
        return ExtractionResult(status=DocumentStatus.UNREADABLE, error=f"openpyxl not installed: {exc}")

    try:
        workbook = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    except Exception as exc:
        logger.warning("Failed to open XLSX %s: %s", path, exc)
        return ExtractionResult(status=DocumentStatus.UNREADABLE, error=str(exc))

    lines: list[str] = []
    for sheet in workbook.worksheets:
        lines.append(f"# Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            if any(cell is not None for cell in row):
                lines.append(" | ".join("" if c is None else str(c) for c in row))

    full_text = "\n".join(lines)
    return ExtractionResult(status=DocumentStatus.OK, text=full_text, page_count=len(workbook.worksheets))


def extract_image(path: Path, ocr_min_confidence: int) -> ExtractionResult:
    try:
        import pytesseract
        from PIL import Image
    except ImportError as exc:
        return ExtractionResult(status=DocumentStatus.UNREADABLE, error=f"pytesseract/Pillow not installed: {exc}")

    try:
        image = Image.open(path)
        text, confidence = _ocr_image(image)
    except Exception as exc:
        logger.warning("Failed to OCR image %s: %s", path, exc)
        return ExtractionResult(status=DocumentStatus.UNREADABLE, error=str(exc))

    status = DocumentStatus.OK
    if confidence is not None and confidence < ocr_min_confidence:
        status = DocumentStatus.OCR_LOW_CONFIDENCE

    return ExtractionResult(
        status=status, text=text.strip(), page_count=1, ocr_used=True, ocr_confidence=confidence
    )


def extract(path: Path, file_type: str, ocr_enabled: bool, ocr_min_confidence: int) -> ExtractionResult:
    """Dispatch to the right extractor by file_type ("pdf" | "docx" | "xlsx" | "image")."""
    if file_type == "pdf":
        return extract_pdf(path, ocr_enabled=ocr_enabled, ocr_min_confidence=ocr_min_confidence)
    if file_type == "docx":
        return extract_docx(path)
    if file_type == "xlsx":
        return extract_xlsx(path)
    if file_type == "image":
        return extract_image(path, ocr_min_confidence=ocr_min_confidence)
    return ExtractionResult(status=DocumentStatus.UNREADABLE, error=f"Unsupported file_type: {file_type}")
